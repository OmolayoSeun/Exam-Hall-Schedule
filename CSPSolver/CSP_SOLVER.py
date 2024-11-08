from resources import Variables as v
from constraint import *
import copy
from docx import Document

# def open_file():
#     with open("", "r") as file:
#         v.courseListJson = json.load(file)

v.courseListJson = {
    "Chemistry": {
        "100 level": {
            "CHM112": [80, 3, "General"],
            "CHM111": [80, 3, "None"],
            "MTH112": [80, 3, "None"],
            "MTH111": [80, 3, "General"],
            "BIO111": [80, 3, "General"],
            "CSC111": [80, 3, "General"],
            "PHY111": [80, 3, "General"],
            "GSE111": [80, 3, "General"],
            "GSE112": [80, 3, "General"]
        },
        "200 level": {
            "CHM211": [80, 3, "None"],
            "CHM212": [80, 3, "None"],
            "CHM213": [80, 3, "General"],
            "CHM214": [80, 3, "None"],
            "CHM219": [80, 3, "None"],
            "PHY211": [80, 3, "None"],
            "MTH211": [80, 3, "None"],
            "MTH213": [80, 3, "None"],
            "GSE211": [80, 3, "General"],
            "CSC211": [80, 3, "None"],
            "EVS212": [80, 3, "None"]
        },
        "300 level": {
            "CHM310": [80, 3, "None"],
            "CHM311": [80, 3, "None"],
            "CHM312": [80, 3, "None"],
            "CHM313": [80, 3, "None"],
            "CHM316": [80, 3, "None"],
            "CHM318": [80, 3, "None"],
            "CHM319": [80, 3, "None"],
            "CHM330": [80, 3, "None"],
            "CHM334": [80, 3, "None"],
            "GSE311": [80, 3, "General"]
        }
    },
    "Geology": {
        "100 level": {
            "GLY111": [80, 3, "None"],
            "BIO111": [80, 3, "General"],
            "CHM112": [80, 3, "General"],
            "CHM119": [80, 3, "None"],
            "PHY111": [80, 3, "General"],
            "MTH111": [80, 3, "General"],
            "CSC111": [80, 3, "General"],
            "GSE111": [80, 3, "General"],
            "GSE112": [80, 3, "General"]
        },
        "200 level": {
            "GLY211": [80, 3, "None"],
            "GLY212": [80, 3, "None"],
            "GLY213": [80, 3, "None"],
            "GLY214": [80, 3, "None"],
            "GLY215": [80, 3, "None"],
            "GLY216": [80, 3, "None"],
            "CHM213": [80, 3, "General"],
            "GSE211": [80, 3, "General"],
            "GSE111": [80, 3, "General"],
            "GSE112": [80, 3, "General"]
        },
        "300 level": {
            "GLY331": [80, 3, "None"],
            "GLY312": [80, 3, "None"],
            "GLY313": [80, 3, "None"],
            "GLY314": [80, 3, "None"],
            "GLY315": [80, 3, "None"],
            "GLY316": [80, 3, "None"],
            "GPH314": [80, 3, "EVS212"],
            "GSE311": [80, 3, "General"]
        }
    }
}
v.hallListJson = {
    "TLH1": 50,
    "TLH2": 50,
    "TLH3": 50,
    "TLH4": 50,
    "TLH5": 50,
    "TLH6": 50,
    "TLH7": 50,
    "TLH8": 50,
    "TLH9": 50,
    "LT1": 250,
    "LT2": 75,
    "LT3": 75,
    "200LT": 100,
    "500LT1": 250,
    "Twin1": 100,
    "Twin2": 100
}
v.availableSlotJson = {
    "Days": 20,
    "Slot Per Day": 3
}
exams = []
examsStudentSize = []

deptExams = []
levelExams = []
levelExamCourseUnit = []
aliasExams = []

examDomain = []
examSolution = {}


def checkForAlias(a, b):
    if a is None:
        return [False, 0]
    else:
        for i in range(len(a)):
            if b in a[i]:
                return [True, i]
        else:
            return [False, 0]


def refineData():
    for level in v.courseListJson.values():
        deptList = []
        for courses in level.values():
            lvlList = []
            lvlListUnit = []
            for courseCode, courseInfo in courses.items():
                if courseInfo[2] == "General":
                    result = checkForAlias(aliasExams, courseCode)
                    if result[0]:
                        courseCode = copy.deepcopy(aliasExams[result[1]][-1] + 'x')
                        aliasExams[result[1]].append(courseCode)

                    else:
                        aliasExams.append([courseCode])
                elif courseInfo[2] != "None":
                    result = checkForAlias(aliasExams, courseInfo[2])
                    if result[0]:
                        aliasExams[result[1]].append(courseCode)
                    else:
                        aliasExams.append([courseInfo[2], courseCode])
                exams.append(courseCode)
                deptList.append(courseCode)
                lvlList.append(courseCode)
                examsStudentSize.append(courseInfo[0])
                lvlListUnit.append(courseInfo[1])

            levelExams.append(lvlList)
            levelExamCourseUnit.append(lvlListUnit)
        deptExams.append(deptList)

    # print(exams, "\n\n")
    # print(examsStudentSize, "\n\n")
    # print(deptExams, "\n\n")
    # print(levelExams, "\n\n")
    # print(levelExamCourseUnit, "\n\n")
    # print(aliasExams, "\n\n")

    halls = []
    days = []
    time = []
    timeSlot = []
    examList = []
    for name, capacity in v.hallListJson.items():
        halls.append([name, capacity])

    for day, slot in v.availableSlotJson.items():
        days.append(day)
        time.append(slot)
    for n in range(len(days)):
        temp = []
        for t in time[n]:
            temp.append([days[n], t])
        timeSlot.append(temp)
    global examDomain
    for days in timeSlot:
        for slot in days:
            for hall in halls:
                # print("hall: ", hall)
                temp = copy.deepcopy(slot)
                temp.append(hall[0])
                temp.append(hall[1])
                # print("days: ", days)
                # print("slot: ", slot)
                # print(temp)
                examList.append(temp)
                del temp
        pass
    for exam in examList:
        string = ''
        for i in range(4):
            string = string + copy.deepcopy(str(exam[i])) + ';'
        examDomain.append(string[:-1])
        del string
    # print(examDomain, end='\n\n')

    # print(len(examDomain))

    return True


def solveCSPProblem():
    problem = Problem()
    for exam in exams:
        problem.addVariable(exam, examDomain)
    problem.addConstraint(AllDifferentConstraint(), exams)
    solutions = problem.getSolution()
    global examSolution
    examSolution = solutions
    # print(solutions)

    return True


def formatToDocument():
    global examSolution
    timetable = {}
    for day, time in v.availableSlotJson.items():
        temp = []
        for t in time:
            temp1 = []
            if t not in temp1:
                temp1.append(t)
            for course, slot in examSolution.items():
                if t in slot and day in slot:
                    temp1.append(course)
            temp.append(temp1)
        timetable[day] = temp
    # for days in timetable.items():
    #     print(days, end="\n\n")
    examSolution = copy.deepcopy(timetable)
    print(examSolution)
    pass


def convertToDocx():
    doc = Document()
    doc.add_heading("Exam Time Table")
    firstKey = ''

    for s in v.availableSlotJson.keys():
        firstKey = s
        break
    colNo = len(v.availableSlotJson[firstKey])
    table = doc.add_table(rows=len(v.availableSlotJson) + 1, cols=colNo + 1)

    hdrCells = table.rows[0].cells
    hdrCells[0].text = 'Days'

    for i in range(colNo):
        hdrCells[i + 1].text = v.availableSlotJson[firstKey][i]

    def find(lst, text):
        for index in range(len(lst)):
            if lst[index] == text:
                return index
        return -1
    i = 0
    for k, j in v.availableSlotJson.items():
        rowCells = table.rows[i + 1].cells
        print(k)
        rowCells[0].text = k
        cExamSol = 0
        for n in range(colNo):
            isAvailable = find(examSolution[k][cExamSol], v.availableSlotJson[firstKey][n])

            print('first key: ', examSolution[firstKey][0], end='\n\n')

            print('isAvailable: ', isAvailable, end='\n\n')

            if isAvailable != -1:
                del examSolution[k][cExamSol][0]
                rowCells[n + 1].text = '\n'.join(examSolution[k][cExamSol])
                cExamSol = cExamSol + 1
            pass
        i = i + 1
        pass
    doc.save("timetable.docx")
    pass


refineData()
solveCSPProblem()
formatToDocument()
convertToDocx()
