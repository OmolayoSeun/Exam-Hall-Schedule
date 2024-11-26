import os
from docx import Document


class CreateDocument:
    def __init__(self, solution, daysCount, slotCount):
        self.solution = solution
        self.dayCount = daysCount
        self.slotCount = slotCount
        self.reorganisedList = []
        pass

    def __refine__(self):

        for day in range(self.dayCount):
            newList = []
            for time in range(self.slotCount):

                string = ''
                for item in self.solution:
                    if item[1] == day and item[2] == time:
                        string = string + item[0] + '[' + item[-1] + ']\n'
                        pass
                newList.append(string)
            self.reorganisedList.append(newList)
        print(self.reorganisedList)

    def create(self):
        self.__refine__()
        doc = Document()
        doc.add_heading("Exam Time Table")

        table = doc.add_table(rows=self.dayCount + 1, cols=self.slotCount + 1)

        hdrCells = table.rows[0].cells
        hdrCells[0].text = 'Days'

        for i in range(self.slotCount):
            hdrCells[i + 1].text = 'time' + str(i + 1)

        for i in range(self.dayCount):
            rowCells = table.rows[i + 1].cells
            rowCells[0].text = 'Day ' + str(i + 1)
            for j in range(self.slotCount):
                rowCells[j + 1].text = self.reorganisedList[i][j]

        filename = "timetable.docx"
        count = 1
        while True:
            if os.path.exists(filename):
                name, extension = filename.rsplit('.', 1)
                filename = f"{name}{count}.{extension}"
            else:
                break

        doc.save(filename)
